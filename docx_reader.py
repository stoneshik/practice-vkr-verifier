"""
Читает DOCX на уровне XML: извлекает сырые абзацы, оглавление и служебные данные документа
"""
import mimetypes
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml.etree import QName

from .constants import (
    ALLOWED_EXTENSIONS,
    ALLOWED_MIME_TYPES,
    PAGE_BREAK_RE,
    OMML_RE,
    PAGE_FIELD_RE,
    FIELD_PAGE_RE,
    FIGURE_CAPTION_RE,
    TABLE_CAPTION_RE,
    XML_NS,
    TOC_LINE_RE
)
from .utils import (
    normalize_text,
    is_list_like,
    list_marker_type,
    is_heading_like,
    is_uppercase,
    get_alignment_name,
    length_to_mm,
    compact_upper
)

def validate_input_file(path: Path) -> None:
    """Проверяет, что входной файл существует и является корректным DOCX"""
    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {path}")
    if not path.is_file():
        raise ValueError(f"Это не файл: {path}")
    if path.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Недопустимое расширение: {path.suffix}")
    guessed_mime, _ = mimetypes.guess_type(str(path))
    if guessed_mime not in ALLOWED_MIME_TYPES:
        raise ValueError(f"Недопустимый MIME-тип: {guessed_mime}")
    if not zipfile.is_zipfile(path):
        raise ValueError("Файл не является корректным DOCX архивом")
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ValueError("Файл не похож на корректный DOCX документ")


def paragraph_has_page_break(paragraph: Paragraph) -> bool:
    """Проверяет наличие разрыва страницы в XML абзаце"""
    try:
        return bool(PAGE_BREAK_RE.search(paragraph._p.xml))
    except Exception:
        return False


def paragraph_has_rendered_page_break(paragraph: Paragraph) -> bool:
    """Проверяет наличие разрыва страницы у абзаца"""
    try:
        xml = paragraph._p.xml
        return "<w:lastRenderedPageBreak" in xml or bool(PAGE_BREAK_RE.search(xml))
    except Exception:
        return False


def paragraph_has_omml(paragraph: Paragraph) -> bool:
    """Проверяет, содержит ли абзац формулу OMML"""
    try:
        return bool(OMML_RE.search(paragraph._p.xml))
    except Exception:
        return False


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    """Проверяет, содержит ли абзац рисунок или графический объект"""
    try:
        xml = paragraph._p.xml
        return "<w:drawing" in xml or "<w:pict" in xml
    except Exception:
        return False


def paragraph_has_numbering(paragraph: Paragraph) -> bool:
    """Проверяет, привязан ли абзац к нумерованному списку Word"""
    try:
        return paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    except Exception:
        return False


def paragraph_numbering_info(paragraph: Paragraph) -> Dict[str, Optional[int]]:
    """Извлекает numId и ilvl для нумерованного абзаца"""
    try:
        ppr = paragraph._p.pPr
        if ppr is None or ppr.numPr is None:
            return {"numId": None, "ilvl": None}
        num_id = None
        ilvl = None
        try:
            if ppr.numPr.numId is not None:
                num_id = int(ppr.numPr.numId.val)
        except Exception:
            pass
        try:
            if ppr.numPr.ilvl is not None:
                ilvl = int(ppr.numPr.ilvl.val)
        except Exception:
            pass
        return {"numId": num_id, "ilvl": ilvl}
    except Exception:
        return {"numId": None, "ilvl": None}


def _resolve_style_chain(style: Any):
    """Итерирует текущий стиль и всю его цепочку базовых стилей"""
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        yield style
        style = getattr(style, "base_style", None)


def resolve_paragraph_format_attr(paragraph: Paragraph, attr: str) -> Any:
    """Ищет атрибут форматирования абзаца напрямую и по цепочке стилей"""
    try:
        value = getattr(paragraph.paragraph_format, attr, None)
        if value is not None:
            return value
    except Exception:
        pass
    style = getattr(paragraph, "style", None)
    for st in _resolve_style_chain(style):
        try:
            pf = getattr(st, "paragraph_format", None)
            if pf is None:
                continue
            value = getattr(pf, attr, None)
            if value is not None:
                return value
        except Exception:
            continue
    return None


def resolve_style_font_attr(paragraph: Paragraph, attr: str) -> Any:
    """Ищет атрибут шрифта абзаца по цепочке стилей"""
    style = getattr(paragraph, "style", None)
    for st in _resolve_style_chain(style):
        try:
            font = getattr(st, "font", None)
            if font is None:
                continue
            value = getattr(font, attr, None)
            if value is not None:
                return value
        except Exception:
            continue
    return None


def iter_block_items(parent: Any) -> Iterable[Any]:
    """Возвращает абзацы и таблицы документа в исходном порядке"""
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _xml_has_page_field(xml_text: str) -> bool:
    """Проверяет, есть ли в XML поле PAGE"""
    return bool(PAGE_FIELD_RE.search(xml_text) or FIELD_PAGE_RE.search(xml_text))


def run_props_to_dict(paragraph: Paragraph) -> Dict[str, Any]:
    """Собирает свойства абзаца в словарь для дальнейших проверок"""
    props = {
        "text": normalize_text(paragraph.text),
        "raw_text": paragraph.text,
        "style": None,
        "alignment": None,
        "font_name": None,
        "font_size_pt": None,
        "bold": None,
        "italic": None,
        "color": None,
        "line_spacing": None,
        "first_line_indent_mm": None,
        "left_indent_mm": None,
        "space_before_pt": None,
        "space_after_pt": None,
        "page_break_before": None,
        "keep_together": None,
        "keep_with_next": None,
        "widow_control": None,
        "is_numbered": paragraph_has_numbering(paragraph),
        "numbering": paragraph_numbering_info(paragraph),
        "has_page_break": paragraph_has_page_break(paragraph),
        "has_rendered_page_break": paragraph_has_rendered_page_break(paragraph),
        "has_drawing": paragraph_has_drawing(paragraph),
        "has_omml": paragraph_has_omml(paragraph),
        "is_list_like": is_list_like(paragraph.text),
        "list_marker_type": list_marker_type(paragraph.text),
        "is_heading_like": is_heading_like(paragraph.text),
        "is_caption": False,
        "is_uppercase": is_uppercase(paragraph.text),
        "contains_abbreviation": False,
    }
    props["is_caption"] = bool(FIGURE_CAPTION_RE.match(props["text"]) or TABLE_CAPTION_RE.match(props["text"]))

    try:
        props["style"] = paragraph.style.name if paragraph.style else None
    except Exception:
        pass
    try:
        props["alignment"] = get_alignment_name(resolve_paragraph_format_attr(paragraph, "alignment"))
    except Exception:
        pass
    try:
        value = resolve_paragraph_format_attr(paragraph, "line_spacing")
        if value is not None:
            props["line_spacing"] = float(value)
    except Exception:
        pass
    try:
        props["first_line_indent_mm"] = length_to_mm(resolve_paragraph_format_attr(paragraph, "first_line_indent"))
    except Exception:
        pass
    try:
        props["left_indent_mm"] = length_to_mm(resolve_paragraph_format_attr(paragraph, "left_indent"))
    except Exception:
        pass
    try:
        sb = resolve_paragraph_format_attr(paragraph, "space_before")
        props["space_before_pt"] = float(sb.pt) if sb is not None else None
    except Exception:
        pass
    try:
        sa = resolve_paragraph_format_attr(paragraph, "space_after")
        props["space_after_pt"] = float(sa.pt) if sa is not None else None
    except Exception:
        pass
    try:
        pbb = resolve_paragraph_format_attr(paragraph, "page_break_before")
        props["page_break_before"] = bool(pbb) if pbb is not None else None
    except Exception:
        pass
    try:
        kt = resolve_paragraph_format_attr(paragraph, "keep_together")
        props["keep_together"] = bool(kt) if kt is not None else None
    except Exception:
        pass
    try:
        kwn = resolve_paragraph_format_attr(paragraph, "keep_with_next")
        props["keep_with_next"] = bool(kwn) if kwn is not None else None
    except Exception:
        pass
    try:
        wc = resolve_paragraph_format_attr(paragraph, "widow_control")
        props["widow_control"] = bool(wc) if wc is not None else None
    except Exception:
        pass

    run_font_names: Counter[str] = Counter()
    run_font_sizes: Counter[float] = Counter()
    run_bold = Counter()
    run_italic = Counter()
    run_colors = Counter()

    for run in paragraph.runs:
        if not normalize_text(run.text):
            continue
        try:
            if run.font.name:
                run_font_names[run.font.name] += 1
        except Exception:
            pass
        try:
            if run.font.size is not None:
                run_font_sizes[round(float(run.font.size.pt), 2)] += 1
        except Exception:
            pass
        try:
            if run.bold is not None:
                run_bold[str(bool(run.bold))] += 1
        except Exception:
            pass
        try:
            if run.italic is not None:
                run_italic[str(bool(run.italic))] += 1
        except Exception:
            pass
        try:
            if run.font.color and run.font.color.rgb:
                run_colors[str(run.font.color.rgb)] += 1
        except Exception:
            pass

    if run_font_names:
        props["font_name"] = run_font_names.most_common(1)[0][0]
    else:
        style_font_name = resolve_style_font_attr(paragraph, "name")
        if style_font_name:
            props["font_name"] = style_font_name

    if run_font_sizes:
        props["font_size_pt"] = run_font_sizes.most_common(1)[0][0]
    else:
        size = resolve_style_font_attr(paragraph, "size")
        if size is not None:
            try:
                props["font_size_pt"] = round(float(size.pt), 2)
            except Exception:
                pass

    if run_bold:
        props["bold"] = run_bold.most_common(1)[0][0] == "True"
    else:
        style_bold = resolve_style_font_attr(paragraph, "bold")
        if style_bold is not None:
            props["bold"] = bool(style_bold)

    if run_italic:
        props["italic"] = run_italic.most_common(1)[0][0] == "True"
    else:
        style_italic = resolve_style_font_attr(paragraph, "italic")
        if style_italic is not None:
            props["italic"] = bool(style_italic)

    if run_colors:
        props["color"] = run_colors.most_common(1)[0][0]

    if not props["is_uppercase"]:
        props["contains_abbreviation"] = bool(
            re.search(r"\b(?:[A-ZА-ЯЁ]{2,5}|[A-ZА-ЯЁ]{1,3}/[A-ZА-ЯЁ]{1,3})\b", props["text"])
        )
    return props


def section_to_dict(section: Any, index: int) -> Dict[str, Any]:
    """Преобразует секцию документа в словарь параметров страницы и колонтитулов"""
    footer_xml = ""
    header_xml = ""
    try:
        footer_xml = section.footer._element.xml
    except Exception:
        pass
    try:
        header_xml = section.header._element.xml
    except Exception:
        pass

    def footer_alignment_name() -> Optional[str]:
        """Определяет преобладающее выравнивание текста в нижнем колонтитуле"""
        try:
            paras = section.footer.paragraphs
            vals = Counter(get_alignment_name(p.alignment) for p in paras if normalize_text(p.text))
            return vals.most_common(1)[0][0] if vals else None
        except Exception:
            return None

    return {
        "section_index": index,
        "left_mm": length_to_mm(section.left_margin),
        "right_mm": length_to_mm(section.right_margin),
        "top_mm": length_to_mm(section.top_margin),
        "bottom_mm": length_to_mm(section.bottom_margin),
        "page_width_mm": length_to_mm(section.page_width),
        "page_height_mm": length_to_mm(section.page_height),
        "different_first_page": bool(getattr(section, "different_first_page_header_footer", False)),
        "odd_and_even_pages": bool(getattr(section, "odd_and_even_pages_header_footer", False)),
        "footer_alignment": footer_alignment_name(),
        "footer_has_page_field": _xml_has_page_field(footer_xml),
        "header_has_page_field": _xml_has_page_field(header_xml),
        "header_has_text": bool(normalize_text(" ".join(p.text for p in section.header.paragraphs))),
        "footer_has_text": bool(normalize_text(" ".join(p.text for p in section.footer.paragraphs))),
        "footer_xml_snippet": footer_xml[:800] if footer_xml else "",
    }


def build_document_model(doc: DocxDocument) -> Dict[str, Any]:
    """Строит внутреннюю модель документа из абзацев, таблиц и секций"""
    blocks: List[Dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            data = run_props_to_dict(block)
            data.update({"kind": "paragraph", "paragraph_index": paragraph_index, "block_index": len(blocks)})
            blocks.append(data)
            paragraph_index += 1
        elif isinstance(block, Table):
            rows = len(block.rows)
            cols = max((len(r.cells) for r in block.rows), default=0)
            has_drawings = False
            text_rows = []
            for row in block.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(normalize_text(cell.text))
                    for p in cell.paragraphs:
                        if paragraph_has_drawing(p):
                            has_drawings = True
                text_rows.append(row_text)
            blocks.append({
                "kind": "table",
                "table_index": table_index,
                "block_index": len(blocks),
                "rows": rows,
                "cols": cols,
                "text_rows": text_rows,
                "has_drawings_in_cells": has_drawings,
            })
            table_index += 1

    sections = [section_to_dict(s, i) for i, s in enumerate(doc.sections)]
    return {
        "blocks": blocks,
        "sections": sections,
        "stats": {
            "paragraphs": sum(1 for b in blocks if b["kind"] == "paragraph"),
            "tables": sum(1 for b in blocks if b["kind"] == "table"),
            "inline_shapes": len(doc.inline_shapes),
            "sections": len(doc.sections),
        }
    }


def extract_raw_paragraph_texts(path: Path) -> List[str]:
    """Извлекает сырые тексты абзацев из document.xml, включая содержимое sdt"""
    with zipfile.ZipFile(path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    body = root.find(".//w:body", namespaces=XML_NS)
    if body is None:
        return []
    result: List[str] = []

    def walk(container):
        """Рекурсивно обходит XML-контейнер и собирает тексты абзацев"""
        for child in container:
            tag = QName(child).localname
            if tag == "p":
                result.append(normalize_text("".join(child.xpath(".//w:t/text()", namespaces=XML_NS))))
            elif tag == "sdt":
                sdt = child.find("w:sdtContent", namespaces=XML_NS)
                if sdt is not None:
                    walk(sdt)

    walk(body)
    return result


def _paragraph_text_with_tabs(p_el: Any) -> str:
    """Собирает текст абзаца из XML с сохранением табуляций и переносов"""
    parts: List[str] = []
    for node in p_el.iter():
        local = QName(node).localname
        if local == "t":
            parts.append(node.text or "")
        elif local == "tab":
            parts.append("\t")
        elif local == "br":
            parts.append("\n")
    raw = "".join(parts)
    raw = raw.replace("\u00a0", " ")
    raw = re.sub(r"[ ]+\t", "\t", raw)
    raw = re.sub(r"\t[ ]+", "\t", raw)
    raw = re.sub(r"[ ]{2,}", " ", raw)
    return raw.strip()


def extract_raw_toc_entries(path: Path) -> List[Dict[str, Any]]:
    """Извлекает строки оглавления из стандартного поля содержания Word"""
    with zipfile.ZipFile(path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    body = root.find(".//w:body", namespaces=XML_NS)
    if body is None:
        return []

    def para_text(el) -> str:
        """Возвращает обычный текст абзаца без сохранения табуляции"""
        return normalize_text("".join(el.xpath(".//w:t/text()", namespaces=XML_NS)))

    def tab_meta(el) -> Dict[str, Any]:
        """Извлекает сведения о табуляции и лидере у абзаца оглавления"""
        tabs = el.xpath("./w:pPr/w:tabs/w:tab", namespaces=XML_NS)
        leaders = [t.get(f"{{{XML_NS['w']}}}leader") for t in tabs if t.get(f"{{{XML_NS['w']}}}leader")]
        vals = [t.get(f"{{{XML_NS['w']}}}val") for t in tabs if t.get(f"{{{XML_NS['w']}}}val")]
        return {
            "has_tab_stop": bool(tabs),
            "tab_vals": vals,
            "leader": leaders[0] if leaders else None,
        }

    for child in body:
        if QName(child).localname != "sdt":
            continue
        sdt_content = child.find("w:sdtContent", namespaces=XML_NS)
        if sdt_content is None:
            continue
        paras = sdt_content.findall("./w:p", namespaces=XML_NS)
        texts = [para_text(p) for p in paras]
        if not texts or compact_upper(texts[0]) != "СОДЕРЖАНИЕ":
            continue

        entries: List[Dict[str, Any]] = []
        for p in paras[1:]:
            raw = _paragraph_text_with_tabs(p)
            line = normalize_text(raw.replace("\t", " "))
            if not line:
                continue

            m = TOC_LINE_RE.match(raw) or TOC_LINE_RE.match(line) or re.match(r"^(.*?)(\d+)\s*$", line)
            if not m:
                continue

            title = normalize_text(m.group(1))
            page = m.group(2)
            if not title:
                continue

            meta = tab_meta(p)
            has_separator = ("\t" in raw) or bool(re.search(r"\.{2,}", raw))or bool(re.search(r"\s{2,}\d+\s*$", raw))
            leader_type = meta["leader"]
            has_leader = bool(leader_type in {"dot", "middleDot", "hyphen", "underscore"} or re.search(r"\.{2,}", raw))

            entries.append({
                "title": title,
                "page": page,
                "has_separator": has_separator or meta["has_tab_stop"],
                "leader_type": leader_type,
                "has_leader": has_leader,
                "raw": raw,
            })
        return entries
    return []


def build_toc_page_lookup(raw_toc_entries: Sequence[Dict[str, Any]]) -> Dict[str, int]:
    """Строит словарь: заголовок из оглавления -> номер страницы"""
    result: Dict[str, int] = {}
    for e in raw_toc_entries:
        title = compact_upper(e.get("title"))
        page = str(e.get("page", "")).strip()
        if title and page.isdigit():
            result[title] = int(page)
    return result
